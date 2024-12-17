# -*- coding: utf-8 -*-
import io
import os
from flask import Flask, session, render_template, request, redirect, url_for, session, flash, current_app, send_file,jsonify,json, make_response
import psycopg2
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from flask import send_from_directory
from datetime import datetime
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO,StringIO
import base64
from docx import Document
from docx.shared import Pt


app = Flask(__name__)
app.secret_key = 'kffdkkfssnfd'  # Change this to a secure secret key
app.static_folder = 'static'

# Replace the following with your PostgreSQL database configuration
db_params = {
    'dbname': 'coursework',
    'user': 'postgres',
    'password': 'postgres',
    'host': 'localhost',
    'port': '5432',
}

# Create a connection to the database
conn = psycopg2.connect(**db_params)
cursor = conn.cursor()

# Assuming you have a table for note


# Create a table for users
cursor.execute('''
    CREATE TABLE IF NOT EXISTS users (
        id SERIAL PRIMARY KEY,
        first_name VARCHAR(50) NOT NULL,
        email VARCHAR(100) UNIQUE NOT NULL,
        password VARCHAR(150) NOT NULL,
        role VARCHAR(50) DEFAULT 'employee' NOT NULL
    )
''')
conn.commit()

# Function to create a new user
def create_user(first_name, email, password, role='employee'):
    hashed_password = generate_password_hash(password, method='pbkdf2:sha256')

    # Insert user data and get the last inserted user ID
    cursor.execute("INSERT INTO users (first_name, email, password, role) VALUES (%s, %s, %s, %s) RETURNING id",
                   (first_name, email, hashed_password, role))
    user_id = cursor.fetchone()[0]
    conn.commit()

    # Insert corresponding employee_info
    cursor.execute("INSERT INTO employee_info (user_id) VALUES (%s)", (user_id,))
    conn.commit()
    cursor.execute("INSERT INTO salary (user_id) VALUES (%s)", (user_id,))
    conn.commit()
    cursor.execute("INSERT INTO higher_education_info (user_id) VALUES (%s)", (user_id,))
    conn.commit()



def get_user_by_email(email):
    try:
        # Execute the SQL query with the email parameter
        cursor.execute("SELECT * FROM users WHERE email = %s", (email,))
        
        # Fetch the result
        user = cursor.fetchone()
        
        return user
    except Exception as e:
        print(f"Error fetching user by email: {str(e)}")
        return None



@app.route('/')
def index():
    return render_template('start_page.html')

@app.route('/sign-up', methods=['GET', 'POST'])
def sign_up():
    if request.method == 'POST':
        first_name = request.form['first_name']
        email = request.form['email']
        password = request.form['password']
        create_user(first_name, email, password)
        return redirect(url_for('login'))
    return render_template('sign_up.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        
        login_status, user_role = validate_user(email, password)

        if login_status:
            session['email'] = email
            flash('Logged in successfully!', category='success')

            return redirect(url_for('home'))  # Redirect to the home route

        else:
            flash('Invalid email or password', category='error')
            return render_template('login.html', error='Invalid email or password')

    return render_template('login.html', error=None)

def validate_user(email, password):
    cursor.execute("SELECT * FROM users WHERE email = %s", (email,))
    user = cursor.fetchone()

    if user:
        stored_hash = user[3]
        stored_role = user[4]  # Assuming role is in the fifth column (index 4)

        if check_password_hash(stored_hash, password):
            return True, stored_role  # Login successful, return role
        else:
            return False, None  # Password mismatch
    else:
        return False, None  # User not found

@app.route('/logout')
def logout():
    session.pop('email', None)  # Remove the 'email' key from the session
    flash('Logged out successfully!', category='success')
    return redirect(url_for('index'))

UPLOAD_FOLDER = r'C:\Users\Public\coursework\coursework\static'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_leave_requests(user_id):
    cursor.execute("SELECT * FROM public.leave_request WHERE user_id = %s", (user_id,))
    leave_requests = cursor.fetchall()
    return leave_requests

@app.route('/home', methods=['GET', 'POST'])
def home():
    if 'email' in session:
        email = session['email']
        user = get_user_by_email(email)

        if user:
            if request.method == 'POST':
                note_text = request.form.get('note')

                # Check if the file is included in the request
                if 'photo' in request.files:
                    photo = request.files['photo']

                    # Check if the file is allowed
                    if photo and allowed_file(photo.filename):
                        # Save the file to the upload folder
                        filename = secure_filename(photo.filename)
                        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                        photo.save(file_path)

                        # Save the file path to the database
                        cursor.execute("UPDATE public.employee_info SET photo = %s WHERE user_id = %s",
                                      (file_path, user[0]))
                        conn.commit()

                        flash('Photo uploaded successfully!', category='success')
                    else:
                        flash('Invalid file type! Allowed types: png, jpg, jpeg, gif', category='error')

            # Fetch additional information from the employee_info table
            employee_info = get_employee_info(user[0])
            leave_requests = get_leave_requests(user[0])

            user_info = {
                'user_id': user[0],
                'first_name': user[1],
                'email': user[2],
                'role': user[4]
                # Assuming role is in the fifth column (index 4)
            }

            # Format hire date without time
            hire_date = employee_info[8].strftime('%Y-%m-%d') if employee_info[8] else None
            termination_date = employee_info[9].strftime('%Y-%m-%d') if employee_info[9] else None

            employee_info_dict = {
                'Name': employee_info[3],
                'Surname': employee_info[4],
                'Department': get_department_name(employee_info[2]),
                'Adress': employee_info[7],
                'hire_date': hire_date,
                'termination_date': termination_date,
                'position': employee_info[11],
                'salary': employee_info[12]
            }

            template_name = f"{user[4].lower()}_home.html"

            return render_template(template_name, user_info=user_info, employee_info=employee_info_dict, leave_requests=leave_requests)

    flash('Unauthorized access!', category='error')
    return redirect(url_for('login'))

# Function to get user's notes
def get_notes(user_id):
    cursor.execute("SELECT id, note_text FROM public.note WHERE user_id = %s", (user_id,))
    notes = cursor.fetchall()
    return notes

# Function to get additional information from the employee_info table
def get_employee_info(user_id):
    cursor.execute("""
        SELECT e.*, s.amount
        FROM public.employee_info e
        LEFT JOIN public.salary s ON e.user_id = s.user_id
        WHERE e.user_id = %s
    """, (user_id,))
    employee_info = cursor.fetchone()
    return employee_info

def get_department_name(department_id):
    cursor.execute("SELECT name FROM public.department WHERE id = %s", (department_id,))
    department_name = cursor.fetchone()
    return department_name[0] if department_name else None

def get_salary(salary_id):
    cursor.execute("SELECT amount FROM public.salary where user_id = %s", (salary_id,))
    salary_amount = cursor.fetchone()
    return salary_amount[0] if salary_amount else None 


@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/submit-leave-request', methods=['POST'])
def submit_leave_request():
    if 'email' in session:
        email = session['email']
        user = get_user_by_email(email)

        if user and user[4] == 'employee' or 'admin' or 'hr_manager':
            start_date = request.form.get('start_date')
            end_date = request.form.get('end_date')

            # Validate and process the leave request
            if validate_leave_request(start_date, end_date):
                cursor.execute("""
                    INSERT INTO public.leave_request (user_id, start_date, end_date, status)
                    VALUES (%s, %s, %s, %s)
                """, (user[0], start_date, end_date, 'Pending'))
                conn.commit()

                flash('Добавлено!', category='success')
            else:
                flash('Неверный формат дат', category='error')

            return redirect(url_for('home'))

    flash('Unauthorized access!', category='error')
    return redirect(url_for('login'))

def validate_leave_request(start_date, end_date):
    # Add your validation logic here
    # For example, check if the dates are valid and if the end_date is after start_date
    return True

# ...

@app.route('/manage-leave-requests', methods=['GET', 'POST'])
def manage_leave_requests():
    if 'email' in session:
        email = session['email']
        user = get_user_by_email(email)

        if user and user[4] == 'hr_manager':
            if request.method == 'POST':
                # Process HR manager actions (accept/decline) on leave requests
                request_id = request.form.get('request_id')
                action = request.form.get('action')

                if action == 'accept':
                    update_leave_request_status(request_id, 'Accepted')
                elif action == 'decline':
                    update_leave_request_status(request_id, 'Declined')

            # Fetch all leave requests for HR manager's review, with sorting by start_date by default
            sort_by = request.args.get('sort_by', 'start_date')
            leave_requests = get_all_leave_requests(sort_by)

            return render_template('manage_leave_requests.html', leave_requests=leave_requests, sort_by=sort_by)

    flash('Unauthorized access!', category='error')
    return redirect(url_for('login'))


def update_leave_request_status(request_id, status):
    cursor.execute("""
        UPDATE public.leave_request
        SET status = %s
        WHERE id = %s
    """, (status, request_id))
    conn.commit()


# Update your get_all_leave_requests function
def get_all_leave_requests(sort_by='start_date'):
    cursor.execute("""
        SELECT lr.id, e."Surname", e."Name", lr.start_date, lr.end_date, lr.status
        FROM leave_request lr
        JOIN employee_info e ON lr.user_id = e.user_id
        ORDER BY
            CASE
                WHEN %s = 'start_date' THEN lr.start_date::text
                WHEN %s = 'end_date' THEN lr.end_date::text
                WHEN %s = 'status' THEN lr.status
                ELSE lr.start_date::text  -- Default sorting criteria
            END;
    """, (sort_by, sort_by, sort_by))
    leave_requests = cursor.fetchall()
    return leave_requests
# Add this to your app.py file
@app.route('/manage-employees', methods=['GET', 'POST'])
def manage_employees():
    if 'email' in session:
        email = session['email']
        user = get_user_by_email(email)

        if user and user[4] == 'hr_manager':
            if request.method == 'POST':
                # Process HR manager actions (update employee info)
                employee_id = request.form.get('employee_id')
                name = request.form.get('name')
                surname = request.form.get('surname')
                department = request.form.get('department')
                salary = request.form.get('salary')
                address = request.form.get('address')
                hire_date = request.form.get('hire_date')
                termination_date = request.form.get('termination_date')
                position = request.form.get('position')

                # Update the employee information
                update_employee_info(employee_id, name, surname, department, salary, address, hire_date, termination_date, position)

            # Fetch all employees for HR manager's review
            employees = get_all_employees_info()
            all_employees = get_all_employees_info()
            search_query = request.form.get('search_query', '').strip().lower()
            if search_query:
                # Filter employees based on the search query
                filtered_employees = [employee for employee in all_employees if search_query in str(employee).lower()]
                return render_template('manage_employees.html', employees=filtered_employees, search_query=search_query)
            else:
                # Render all employees if no search query is provided
                return render_template('manage_employees.html', employees=all_employees)
            
        return render_template('manage_employees.html', employees=employees)

    flash('Unauthorized access!', category='error')
    return redirect(url_for('login'))

# Add these functions to your database module (e.g., db.py)
def get_all_employees_info(sort_by='hire_date'):
    # Validate sort_by parameter to prevent SQL injection
    allowed_sort_columns = ['id', 'hire_date', 'name', 'surname', 'amount']
    if sort_by not in allowed_sort_columns:
        sort_by = 'hire_date'  # Default sorting column

    cursor.execute("""
        SELECT e.user_id, e."Name", e."Surname", d."name" as Department, s.amount AS "Salary", e."Adress", e.hire_date, e.termination_date, e.position, ed.school, ed.degree, ed.major , ed.graduation_date
        FROM employee_info e
        LEFT JOIN salary s ON e.user_id = s.user_id
        LEFT JOIN department d on e.department_id = d.id
		LEFT JOIN higher_education_info ed on e.user_id = ed.user_id;
       """)

    employees = cursor.fetchall()
    return employees

def get_employee_info_by_id(employee_id):
    cursor.execute("""
        SELECT e.user_id, e."Name", e."Surname", d."name" as Department, s.amount AS "Salary", e."Adress", e.hire_date, e.termination_date, e."position", ed.school, ed.degree, ed.major , ed.graduation_date
        FROM employee_info e
        LEFT JOIN salary s ON e.user_id = s.user_id
        LEFT JOIN department d on e.department_id = d.id
        LEFT JOIN higher_education_info ed on e.user_id = ed.user_id             
        WHERE e.user_id = %s
    """, (employee_id,))
    employee_info = cursor.fetchone()
    return employee_info



def update_employee_info(employee_id, name, surname, department_id, salary, address, hire_date, termination_date, position):
    # Check if employee_id is not empty
    if not employee_id:
        flash('Invalid employee ID!', category='error')
        return

    try:
        # Update the employee_info table
        cursor.execute("""
            UPDATE public.employee_info
            SET "Name" = %s, "Surname" = %s, "Adress" = %s, "department_id" = %s, "hire_date" = %s, "termination_date" = %s, "position" = %s
            WHERE user_id = %s
        """, (name, surname, address,department_id, hire_date, termination_date, position, employee_id))
        
        if salary != '':
            # Update the salary table
            cursor.execute("""
                UPDATE public.salary
                SET amount = %s
                WHERE user_id = %s
            """, (salary, employee_id))

        # Commit the transaction
        conn.commit()
        flash('Employee information updated successfully!', category='success')
    except Exception as e:
        # Rollback the transaction in case of any error
        conn.rollback()
        flash(f'Error updating employee information: {str(e)}', category='error')



@app.route('/employee/<int:employee_id>', methods=['GET', 'POST'])
def employee_page(employee_id):
    if 'email' in session:
        email = session['email']
        user = get_user_by_email(email)

        if user and user[4] == 'hr_manager':
            if request.method == 'POST':
                # Process HR manager actions (update employee info)
                name = request.form.get('name')
                surname = request.form.get('surname')
                department_id = request.form.get('department')
                salary = request.form.get('salary')
                address = request.form.get('address')
                hire_date = request.form.get('hire_date')
                termination_date = request.form.get('termination_date')
                position = request.form.get('position')

                # Update the employee information
                update_employee_info(employee_id, name, surname, department_id, salary, address, hire_date, termination_date, position)

            # Fetch the information for the selected employee
            employee_info = get_employee_info_by_id(employee_id)                  
            if 'delete_employee' in request.form:

            # Handle employee deletion (implement this function)
                delete_employee(employee_id)
                flash('Employee deleted successfully!', category='success')
                return redirect(url_for('manage_employees'))
            else:
                return render_template('employee_page.html', employee_info=employee_info)
                   

        
    flash('Unauthorized access!', category='error')
    return redirect(url_for('login'))

def delete_employee(employee_id):
    try:
        # Delete from the 'employee_info' table
        cursor.execute("DELETE FROM employee_info WHERE user_id = %s", (employee_id,))
        conn.commit()

        # Delete from the 'users' table (assuming 'id' is the primary key in 'users' table)
        cursor.execute("DELETE FROM users WHERE id = %s", (employee_id,))
        conn.commit()

        # Delete from the 'salary' table (adjust table/column names accordingly)
        cursor.execute("DELETE FROM higher_education_info WHERE user_id = %s", (employee_id,))
        conn.commit()

        cursor.execute("DELETE FROM salary WHERE user_id = %s", (employee_id,))
        conn.commit()

        # Add more delete statements for other related tables

        flash('Employee deleted successfully!', category='success')

    except Exception as e:
        print(f"Error deleting employee: {str(e)}")
        conn.rollback()
        flash('Error deleting employee!', category='error')


@app.route('/manage-employees-education', methods=['GET', 'POST'])
def manage_employees_education():
    if 'email' in session:
        email = session['email']
        user = get_user_by_email(email)

        if user and user[4] == 'hr_manager':
            if request.method == 'POST':
                # Process HR manager actions (update employee info)
                employee_id = request.form.get('employee_id')
                name = request.form.get('name')
                surname = request.form.get('surname')
                uni = request.form.get('school')
                degree = request.form.get('degree')
                major = request.form.get('major')
                grad_date = request.form.get('graduation_date')
                # Update the employee information
                update_employee_education_info(employee_id,uni,degree,major,grad_date)

            # Fetch all employees for HR manager's review
            employees = get_all_employees_info()
            all_employees = get_all_employees_info()
            search_query = request.form.get('search_query', '').strip().lower()
            if search_query:
                # Filter employees based on the search query
                filtered_employees = [employee for employee in all_employees if search_query in str(employee).lower()]
                return render_template('manage_employees_education.html', employees=filtered_employees, search_query=search_query)
            else:
                # Render all employees if no search query is provided
                return render_template('manage_employees_education.html', employees=all_employees)
            
        return render_template('manage_employees_education.html', employees=employees)

    flash('Unauthorized access!', category='error')
    return redirect(url_for('login'))

@app.route('/employee-ed/<int:employee_id>', methods=['GET', 'POST'])
def employee_page_ed(employee_id):
    if 'email' in session:
        email = session['email']
        user = get_user_by_email(email)

        if user and user[4] == 'hr_manager':
            if request.method == 'POST':
                # Process HR manager actions (update employee info)
                employee_id = request.form.get('employee_id')
                name = request.form.get('name')
                surname = request.form.get('surname')
                uni = request.form.get('school')
                degree = request.form.get('degree')
                major = request.form.get('major')
                grad_date = request.form.get('graduation_date')
                # Update the employee information
                update_employee_education_info(employee_id,uni,degree,major,grad_date)

            # Fetch the information for the selected employee
            employee_info = get_employee_info_by_id(employee_id)

            return render_template('employee_page_ed.html', employee_info=employee_info)

    flash('Unauthorized access!', category='error')
    return redirect(url_for('login'))

# Add these functions to your database module (e.g., db.py)

def update_employee_education_info(employee_id,uni,degree,major,grad_date):
    # Check if employee_id is not empty
    if not employee_id:
        flash('Invalid employee ID!', category='error')
        return

    try:
        # Update the employee_info table
        cursor.execute("""
            UPDATE public.higher_education_info
            SET school = %s, degree = %s, major = %s, graduation_date = %s
            WHERE id = (SELECT user_id FROM public.employee_info WHERE user_id = %s)
        """, (uni, degree, major, grad_date, employee_id))

        conn.commit()
        flash('Employee information updated successfully!', category='success')
    except Exception as e:
        # Rollback the transaction in case of any error
        conn.rollback()
        flash(f'Error updating employee information: {str(e)}', category='error')


# Existing route for downloading PDF
@app.route('/download-pdf/<int:employee_id>', methods=['GET'])
def download_employee_pdf(employee_id):
    employee_info = get_employee_info_by_id(employee_id)

    if employee_info:
        pdf_document = generate_pdf(employee_info)
        return send_file(io.BytesIO(pdf_document), as_attachment=True, download_name='employee_document.pdf', mimetype='application/pdf')
    else:
        flash('Employee not found!', category='error')
        return redirect(url_for('manage_employees'))

# Existing route for downloading Excel and CSV
@app.route('/download-excel/<file_format>', methods=['GET'])
def download_employee_excel(file_format):
    employees_info = get_all_employees_info()

    if employees_info:
        df = pd.DataFrame(employees_info, columns=['Name', 'Surname', 'Department', 'Salary', 'Address', 'Hire Date', 'Termination Date'])

        if file_format == 'xlsx':
            excel_data = df.to_excel(index=False, sheet_name='Employee Information')
            return send_file(excel_data, as_attachment=True, download_name='employee_information.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        elif file_format == 'csv':
            csv_data = df.to_csv(index=False)
            return send_file(io.BytesIO(csv_data.encode()), as_attachment=True, download_name='employee_information.csv', mimetype='text/csv')
        else:
            flash('Invalid file format!', category='error')
            return redirect(url_for('manage_employees'))
    else:
        flash('No employees found!', category='error')
        return redirect(url_for('manage_employees'))

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_UNDERLINE
# New route for downloading employee information in different formats
@app.route('/download-employee-info/<int:employee_id>/<file_format>', methods=['GET'])
def download_employee_info(employee_id, file_format):
    employee_info = get_employee_info_by_id(employee_id)

    if employee_info:
        if file_format == 'pdf':
            return download_pdf(employee_info)        
        elif file_format == 'docx1':
            return download_document_fire_tenant(employee_info)
        elif file_format == 'docx2':
            return download_document_reward(employee_info)
        elif file_format == 'docx3':
            return download_document_punishment(employee_info) 
        elif file_format == 'docx4':
            return download_document_bussines_trip(employee_info)       
        else:
            flash('Invalid file format!', category='error')
            return redirect(url_for('manage_employees'))
    else:
        flash('Employee not found!', category='error')
        return redirect(url_for('manage_employees'))

# Function to download PDF
def download_pdf(employee_info):
    pdf_document = generate_pdf(employee_info)
    return send_file(io.BytesIO(pdf_document), as_attachment=True, download_name='employee_document.pdf', mimetype='application/pdf')
# Function to download Excel or CSV

# Function to download Word document
def download_document_fire_tenant(employee_info):
    # Создание нового документа
    document = Document()
# Параграф с центрированным текстом в три строки
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    title_text = "ПРИКАЗ ОБ УВОЛЬНЕНИИ РАБОТНИКА ПО СОБСТВЕННОМУ\nЖЕЛАНИЮ В СВЯЗИ С НАРУШЕНИЕМ НАНИМАТЕЛЕМ\nЗАКОНОДАТЕЛЬСТВА О ТРУДЕ"
    paragraph.add_run(title_text)

# Параграф с выровненным по левому краю текстом
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    paragraph.add_run("ПРИКАЗ\n") 
    paragraph.add_run(" ")
    paragraph.add_run("Об увольнении с работы\n(освобождении от занимаемой\nдолжности)")

# Параграф с центрированным текстом, шрифт меньше
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph.add_run("УВОЛИТЬ (ОСВОБОДИТЬ ОТ ЗАНИМАЕМОЙ ДОЛЖНОСТИ):")

# Параграфы с центрированными данными из базы данных
    name = str(employee_info[1]) + " " + str(employee_info[2])
        
    position = str(employee_info[8])
    
    department = str(employee_info[3])
    

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph.add_run(name)
    paragraph = document.add_paragraph()
    run = paragraph.add_run("______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________")
    run.font.size = Pt(2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run("(Ф.И.О)")
    run.font.size = Pt(8)
    
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    paragraph.add_run(position)
    paragraph = document.add_paragraph()
    run1 = paragraph.add_run("______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________")
    run1.font.size = Pt(2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = paragraph.add_run("(название занимаемой должности)")
    run2.font.size = Pt(8)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph.add_run(department)
    paragraph = document.add_paragraph()
    run3 = paragraph.add_run("______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________")
    run3.font.size = Pt(2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run4 = paragraph.add_run("(название отддела)")
    run4.font.size = Pt(8)

# Параграф с центрированным текстом
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph.add_run("По собственному желанию в связи с нарушением нанимателем законодательства о труде (ст. 40 Трудового кодекса Республики Беларусь).")

# Параграф с центрированным текстом
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph.add_run("Бухгалтерии выплатить выходное пособие в размере двухнедельного среднего заработка.")

# Параграф с выровненным по левому краю текстом
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.add_run("Основание:")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.add_run("Руководитель организации: ")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph.add_run("подпись                 расшифровка подписи ")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.add_run("С приказом ознакомлен:")

    paragraph = document.add_paragraph()   
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph.add_run("подпись                 расшифровка подписи ")# Создание временного файлового объекта
    temp_file = io.BytesIO()
    document.save(temp_file)
    temp_file.seek(0)

    # Создание ответа Flask
    response = make_response(temp_file.getvalue())
    response.headers['Content-Disposition'] = 'attachment; filename=employee_information.docx'
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    return response

def download_document_reward(employee_info):
    # Создание нового документа
    document = Document()
# Параграф с центрированным текстом в три строки
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    title_text = "ПРИКАЗ О ПООЩРЕНИИ"
    paragraph.add_run(title_text)
    paragraph = document.add_paragraph()
    run = paragraph.add_run(" ")
    run.font.size = Pt(14)

# Параграф с выровненным по левому краю текстом
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    paragraph.add_run("ПРИКАЗ № \n _____________________") 
    paragraph = document.add_paragraph()
    paragraph.add_run("О поощрении")

# Параграф с центрированным текстом, шрифт меньше
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    run = paragraph.add_run("______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________")
    run.font.size = Pt(2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run("(мотив поощрения)")
    run.font.size = Pt(8)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    run = paragraph.add_run("______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________")
    run.font.size = Pt(2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run("(вид поощрения)")
    run.font.size = Pt(8)

# Параграфы с центрированными данными из базы данных
    name = str(employee_info[1]) + " " + str(employee_info[2])
        
    position = str(employee_info[8])
    
    department = str(employee_info[3])
    

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph.add_run(name)
    paragraph = document.add_paragraph()
    run = paragraph.add_run("______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________")
    run.font.size = Pt(2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run("(Ф.И.О)")
    run.font.size = Pt(8)
    
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    paragraph.add_run(position)
    paragraph = document.add_paragraph()
    run1 = paragraph.add_run("______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________")
    run1.font.size = Pt(2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = paragraph.add_run("(название занимаемой должности)")
    run2.font.size = Pt(8)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph.add_run(department)
    paragraph = document.add_paragraph()
    run3 = paragraph.add_run("______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________")
    run3.font.size = Pt(2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run4 = paragraph.add_run("(название отдела)")
    run4.font.size = Pt(8)

# Параграф с центрированным текстом
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph.add_run("   Основание:_______________________________________________________________")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.add_run("Руководитель организации: ")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph.add_run("подпись                 расшифровка подписи ")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.add_run("С приказом ознакомлен:")

    paragraph = document.add_paragraph()   
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph.add_run("подпись                 расшифровка подписи ")# Создание временного файлового объекта
    temp_file = io.BytesIO()
    document.save(temp_file)
    temp_file.seek(0)

    # Создание ответа Flask
    response = make_response(temp_file.getvalue())
    response.headers['Content-Disposition'] = 'attachment; filename=employee_information.docx'
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    return response

def download_document_punishment(employee_info):
    # Создание нового документа
    document = Document()
# Параграф с центрированным текстом в три строки
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    title_text = "ПРИКАЗ О НАЛОЖЕНИИ ДИСЦИПЛИНАРНОГО ВЗЫСКАНИЯ"
    paragraph.add_run(title_text)
    paragraph = document.add_paragraph()
    run = paragraph.add_run(" ")
    run.font.size = Pt(14)

# Параграф с выровненным по левому краю текстом
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    paragraph.add_run("ПРИКАЗ № \n _____________________") 
    paragraph = document.add_paragraph()
    paragraph.add_run("О наложении дисциплинарного взыскания\n _____________________ ")
    paragraph = document.add_paragraph()
    run = paragraph.add_run("(наименование взыскания)")
    run.font.size = Pt(8)


# Параграф с центрированным текстом, шрифт меньше
    
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


# Параграфы с центрированными данными из базы данных
    name = str(employee_info[1]) + " " + str(employee_info[2])
        
    position = str(employee_info[8])
    
    department = str(employee_info[3])
    

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph.add_run(name)
    paragraph = document.add_paragraph()
    run = paragraph.add_run("______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________")
    run.font.size = Pt(2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run("(Ф.И.О)")
    run.font.size = Pt(8)
    
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    paragraph.add_run(position)
    paragraph = document.add_paragraph()
    run1 = paragraph.add_run("______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________")
    run1.font.size = Pt(2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = paragraph.add_run("(название занимаемой должности)")
    run2.font.size = Pt(8)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph.add_run(department)
    paragraph = document.add_paragraph()
    run3 = paragraph.add_run("______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________")
    run3.font.size = Pt(2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run4 = paragraph.add_run("(название отдела)")
    run4.font.size = Pt(8)

    paragraph = document.add_paragraph()
    run3 = paragraph.add_run("______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________")
    run3.font.size = Pt(2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run4 = paragraph.add_run("(вид нарушения, дата нарушения)")
    run4.font.size = Pt(8)

# Параграф с центрированным текстом
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragraph.add_run("   Основание:_______________________________________________________________")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.add_run("Руководитель организации: ")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph.add_run("подпись                 расшифровка подписи ")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.add_run("С приказом ознакомлен:")

    paragraph = document.add_paragraph()   
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph.add_run("подпись                 расшифровка подписи ")# Создание временного файлового объекта
    temp_file = io.BytesIO()
    document.save(temp_file)
    temp_file.seek(0)

    # Создание ответа Flask
    response = make_response(temp_file.getvalue())
    response.headers['Content-Disposition'] = 'attachment; filename=employee_information.docx'
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    return response

def download_document_bussines_trip(employee_info):
    # Создание нового документа
    document = Document()
# Параграф с центрированным текстом в три строки
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    title_text = "ПРИКАЗ О КОМАНДИРОВАНИИ"
    paragraph.add_run(title_text)
    paragraph = document.add_paragraph()
    run = paragraph.add_run(" ")
    run.font.size = Pt(14)

# Параграф с выровненным по левому краю текстом
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    paragraph.add_run("ПРИКАЗ № \n _____________________") 
    paragraph = document.add_paragraph()
    paragraph.add_run("О командировании\n _____________________ ")
    paragraph = document.add_paragraph()
    


# Параграф с центрированным текстом, шрифт меньше
    
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


# Параграфы с центрированными данными из базы данных
    name = str(employee_info[1]) + " " + str(employee_info[2])
        
    position = str(employee_info[8])
    
    department = str(employee_info[3])
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph.add_run("     В связи с______________________________________________________________ ")
    paragraph = document.add_paragraph()
    run = paragraph.add_run("______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________")
    run.font.size = Pt(2)
    paragraph = document.add_paragraph()
    run = paragraph.add_run("(основание)")
    run.font.size = Pt(8)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.add_run("ПРИКАЗЫВАЮ:  ")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph.add_run("Командировать, ")
    paragraph.add_run(name)
    paragraph = document.add_paragraph()
    run = paragraph.add_run("______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________")
    run.font.size = Pt(2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run("(Ф.И.О)")
    run.font.size = Pt(8)
    
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    paragraph.add_run(position)
    paragraph = document.add_paragraph()
    run1 = paragraph.add_run("______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________")
    run1.font.size = Pt(2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = paragraph.add_run("(название занимаемой должности)")
    run2.font.size = Pt(8)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    paragraph.add_run(department)
    paragraph = document.add_paragraph()
    run3 = paragraph.add_run("______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________")
    run3.font.size = Pt(2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run4 = paragraph.add_run("(название отдела)")
    run4.font.size = Pt(8)

    paragraph = document.add_paragraph()
    paragraph.add_run("")
    paragraph = document.add_paragraph()
    run3 = paragraph.add_run("______________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________________")
    run3.font.size = Pt(2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run4 = paragraph.add_run("(место командирования, даты начала и окончания)")
    run4.font.size = Pt(8)

# Параграф с центрированным текстом
    
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.add_run("Руководитель организации: ")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph.add_run("подпись                 расшифровка подписи ")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph.add_run("С приказом ознакомлен:")

    paragraph = document.add_paragraph()   
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph.add_run("подпись                 расшифровка подписи ")# Создание временного файлового объекта
    temp_file = io.BytesIO()
    document.save(temp_file)
    temp_file.seek(0)

    # Создание ответа Flask
    response = make_response(temp_file.getvalue())
    response.headers['Content-Disposition'] = 'attachment; filename=employee_information.docx'
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    return response

def generate_pdf(employee_info):
    pdf_buffer = io.BytesIO()
    p = canvas.Canvas(pdf_buffer)
    
    font_name = "times"  
    pdfmetrics.registerFont(TTFont(font_name, r"C:\Users\Public\coursework\coursework\static\times.ttf"))

    # Set the font for the canvas
    p.setFont(font_name, 14)


    # Заголовок приказа
    p.drawString(100, 750, "Докладная записка о поощрении")
    p.drawString(100, 730, "№")
    p.drawString(220, 730, "Место издания")
    p.drawString(100, 710, "О предоставлении отпуска")
    # Данные о сотруднике
    p.drawString(100, 680, "ПРЕДОСТАВИТЬ:")
    p.drawString(100, 660, f"ФАМИЛИЯ, имя, отчество работника: {employee_info[1]} {employee_info[2]}")
    p.drawString(100, 640, f"Наименование должности: {employee_info[3]}")
    p.drawString(100, 620, f"Наименование структурного подразделения: {employee_info[4]}")
    p.drawString(100, 600, f"Вид отпуска: {employee_info[5]}")
    p.drawString(100, 580, f"Количество календарных дней отпуска: {employee_info[6]}")
    p.drawString(100, 560, f"с {employee_info[7]} по {employee_info[8]}")
    p.drawString(100, 540, f"Основание:")

    # Нижняя часть приказа
    p.drawString(100, 480, "Наименование должности руководителя организации")
    p.drawString(100, 460, "Подпись")
    p.drawString(100, 440, "Расшифровка подписи")

    p.drawString(100, 400, "Визы")

    p.drawString(100, 360, "С приказом ознакомлен")
    p.drawString(100, 340, "Подпись")
    p.drawString(100, 320, "Расшифровка подписи")
    p.drawString(100, 300, "Дата")

    p.showPage()
    p.save()

    return pdf_buffer.getvalue()

def generate_pie_chart(data, title):
    labels = data.keys()
    sizes = data.values()

    plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
    plt.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.

    plt.title(title)

    # Save the plot to a BytesIO buffer
    image_buffer = BytesIO()
    plt.savefig(image_buffer, format='png')
    plt.close()

    # Convert the plot to base64 for embedding in HTML
    image_base64 = base64.b64encode(image_buffer.getvalue()).decode('utf-8')

    return image_base64

@app.route('/pie-charts', methods=['GET'])
def pie_charts():
    try:
        # Default values for charts
        majors_chart = None
        degrees_chart = None

        # Fetch data for majors and degrees
        majors_data = get_majors_data()
        degrees_data = get_degrees_data()

        # Check if data is not None before generating charts
        if majors_data is not None and degrees_data is not None:
            # Generate pie chart images
            majors_chart = generate_pie_chart(majors_data, 'Distribution of Majors')
            degrees_chart = generate_pie_chart(degrees_data, 'Distribution of Degrees')

        return render_template('pie_charts.html', majors_chart=majors_chart, degrees_chart=degrees_chart, user_info=session.get('user_info'))
    except Exception as e:
        flash(f'An error occurred: {str(e)}', category='error')
        return render_template('pie_charts.html', majors_chart=majors_chart, degrees_chart=degrees_chart, user_info=session.get('user_info'))


@app.route('/projects', methods=['GET'])
def projects():
    try:
        # Default values for charts
        majors_chart = None
        degrees_chart = None

        # Fetch data for majors and degrees
        majors_data = get_majors_data()
        degrees_data = get_degrees_data()

        # Check if data is not None before generating charts
        if majors_data is not None and degrees_data is not None:
            # Generate pie chart images
            majors_chart = generate_pie_chart(majors_data, 'Distribution of Majors')
            degrees_chart = generate_pie_chart(degrees_data, 'Distribution of Degrees')

        return render_template('pie_charts.html', majors_chart=majors_chart, degrees_chart=degrees_chart, user_info=session.get('user_info'))
    except Exception as e:
        flash(f'An error occurred: {str(e)}', category='error')
        return render_template('pie_charts.html', majors_chart=majors_chart, degrees_chart=degrees_chart, user_info=session.get('user_info'))

def get_majors_data():
    try:
        cursor.execute("SELECT major, COUNT(*) FROM higher_education_info GROUP BY major")
        majors_data = dict(cursor.fetchall())
        return majors_data
    except Exception as e:
        print(f"Error fetching majors data: {str(e)}")
        return None

def get_degrees_data():
    try:
        cursor.execute("SELECT degree, COUNT(*) FROM higher_education_info GROUP BY degree")
        degrees_data = dict(cursor.fetchall())
        return degrees_data
    except Exception as e:
        print(f"Error fetching degrees data: {str(e)}")
        return None
    
def get_projects_data():
    try:
        cursor.execute("SELECT degree, COUNT(*) FROM higher_education_info GROUP BY degree")
        degrees_data = dict(cursor.fetchall())
        return degrees_data
    except Exception as e:
        print(f"Error fetching degrees data: {str(e)}")
        return None

def export_to_csv():
    # Fetch data from the database
    data = fetch_data_from_database()

    # Convert data to a DataFrame
    df = pd.DataFrame(data, columns=['column1', 'column2', ...])

    # Export to CSV
    df.to_csv('exported_data.csv', index=False)

def fetch_data_from_database():
    try:
        # Connect to the database
        conn = psycopg2.connect(
            dbname='your_database_name',
            user='your_username',
            password='your_password',
            host='your_host',
            port='your_port'
        )

        # Create a cursor
        cursor = conn.cursor()

        # Fetch data from the database (replace with your actual query)
        query = "SELECT * FROM your_table"
        cursor.execute(query)
        data = cursor.fetchall()

        # Close the cursor and connection
        cursor.close()
        conn.close()

        return data

    except Exception as e:
        print(f"Error fetching data from the database: {str(e)}")
        return None

def import_from_csv(file_path):
    try:
        # Read data from CSV
        df = pd.read_csv(file_path)

        # Insert data into the database
        insert_data_into_database(df)

        flash('Data imported successfully!', category='success')

    except Exception as e:
        print(f"Error importing data: {str(e)}")
        flash('Error importing data!', category='error')
def insert_data_into_database(df):
    try:
        # Connect to the database
        conn = psycopg2.connect(
            dbname='your_database_name',
            user='your_username',
            password='your_password',
            host='your_host',
            port='your_port'
        )

        # Create a cursor
        cursor = conn.cursor()

        # Convert DataFrame to CSV string
        csv_data = df.to_csv(index=False, header=False, sep=',')

        # Create a temporary CSV file-like object for COPY FROM
        csv_file = StringIO(csv_data)

        # Execute COPY FROM to insert data into the database
        cursor.copy_from(csv_file, 'your_table', sep=',')

        # Commit changes and close the cursor and connection
        conn.commit()
        cursor.close()
        conn.close()

        print("Data inserted successfully.")

    except Exception as e:
        print(f"Error inserting data into the database: {str(e)}")


@app.route('/export-database', methods=['POST'])
def export_database():
    # Implement your logic to fetch data from the database
    # For example, fetch all data from the 'employee_info' table
    cursor.execute("SELECT * FROM employee_info")
    employee_data = cursor.fetchall()

    # Convert the data to a list of dictionaries (each row is a dictionary)
    employee_data_list = []
    for row in employee_data:
        employee_data_list.append({
            'id': row[0],
            'name': row[1],
            'surname': row[2],
            'department': row[3],
            'salary': row[4],
            'address': row[5],
            'hire_date': row[6],
            'termination_date': row[7]
            # Add other fields as needed
        })

    # Create a JSON response
    response = jsonify({'employees': employee_data_list})

    # Set the Content-Disposition header to prompt the user to download the JSON file
    response.headers['Content-Disposition'] = 'attachment; filename=employee_data.json'

    return response

@app.route('/import-database', methods=['POST'])
def import_database():
    try:
        # Get the uploaded file
        uploaded_file = request.files['file']

        # Check if the file is a JSON file
        if uploaded_file and uploaded_file.filename.endswith('.json'):
            # Read the JSON data from the file
            data = json.loads(uploaded_file.read())

            if 'departments' in data and isinstance(data['departments'], list):
                # Assuming data is a list of dictionaries
                for department in data['departments']:
                    if 'id' in department and 'name' in department:
                        cursor.execute("""
                            INSERT INTO department (id, name)
                            VALUES (%s, %s)
                            ON CONFLICT (id) DO UPDATE SET name = EXCLUDED.name""",
                            (department['id'], department['name']))

                        conn.commit()
                    else:
                        flash('Invalid data format! Each department should have "id" and "name" fields.', category='error')
                        return redirect(url_for('home'))
                
                flash('Data imported successfully!', category='success')
            else:
                flash('Invalid data format! The file should contain a "departments" list.', category='error')
        else:
            flash('Invalid file format! Please upload a JSON file.', category='error')

    except Exception as e:
        error_message = f'Error importing data: {str(e)}'
        flash(error_message, category='error')
        print(error_message)
        
    return redirect(url_for('home'))


if __name__ == '__main__':
    app.run(debug=True)
